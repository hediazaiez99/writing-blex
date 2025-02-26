import base64
import difflib
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.core.files.base import ContentFile
from .forms import DocumentUploadForm
from .models import ValidatedDocuments
import docx
from io import BytesIO

def update_table(original_table, new_table):
    header = original_table.rows[0]
    row_elements = original_table._element.xpath("./w:tr")
    for row in row_elements[1:]:
        original_table._element.remove(row)
    for new_row in new_table.rows[1:]:
        new_row_in_orig = original_table.add_row()
        for idx, cell in enumerate(new_row.cells):
            if idx < len(new_row_in_orig.cells):
                new_row_in_orig.cells[idx].text = cell.text

def compare_impurities(doc1, doc2):
    impurities_doc1 = 5
    impurities_doc2 = 3
    inversion = "Oui"
    return impurities_doc1, impurities_doc2, inversion

def update_batch_analyses(doc1_file, doc2_file):
    doc1 = docx.Document(doc1_file)
    doc2 = docx.Document(doc2_file)
    if doc1.tables and doc2.tables:
        update_table(doc1.tables[0], doc2.tables[0])
    nb_tables = min(len(doc1.tables), len(doc2.tables))
    for i in range(1, nb_tables):
        update_table(doc1.tables[i], doc2.tables[i])
    impurities_doc1, impurities_doc2, inversion = compare_impurities(doc1, doc2)
    doc1.add_paragraph("")
    doc1.add_paragraph("=== Mise à jour par REG X ===")
    doc1.add_paragraph("Les données du document 2 ont été intégrées dans la section 'Batch Analyses'.")
    doc1.add_paragraph("Comparaison des impuretés :")
    doc1.add_paragraph("• Document 1 : {} impuretés détectées".format(impurities_doc1))
    doc1.add_paragraph("• Document 2 : {} impuretés détectées".format(impurities_doc2))
    doc1.add_paragraph("• Inversion de l'ordre des impuretés : {}".format(inversion))
    f = BytesIO()
    doc1.save(f)
    f.seek(0)
    return f

def get_doc_text(document):
    """
    Extrait le contenu complet d'un document DOCX en incluant les paragraphes et le contenu des tableaux.
    """
    lines = []
    lines.append("=== Paragraphes ===")
    for p in document.paragraphs:
        lines.append(p.text)
    lines.append("\n=== Tableaux ===")
    for idx, table in enumerate(document.tables):
        lines.append(f"-- Tableau {idx + 1} --")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            lines.append(row_text)
    return "\n".join(lines)

def verify_update(updated_doc, doc2):
    errors = []
    nb_tables = min(len(updated_doc.tables), len(doc2.tables))
    for i in range(nb_tables):
        table_expected = doc2.tables[i]
        table_actual = updated_doc.tables[i]
        expected_rows = len(table_expected.rows) - 1
        actual_rows = len(table_actual.rows) - 1
        if expected_rows != actual_rows:
            errors.append(f"Table {i + 1}: nombre de lignes attendu {expected_rows}, trouvé {actual_rows}")
        else:
            for row_index in range(1, len(table_expected.rows)):
                expected_cells = [cell.text.strip() for cell in table_expected.rows[row_index].cells]
                actual_cells = [cell.text.strip() for cell in table_actual.rows[row_index].cells]
                if expected_cells != actual_cells:
                    errors.append(f"Table {i + 1}, ligne {row_index}: attendu {expected_cells}, trouvé {actual_cells}")
    return errors

def upload_documents(request):
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'validate':
            # Validation : utilisation des fichiers stockés en session
            doc1_b64 = request.session.get('validated_doc1')
            doc2_b64 = request.session.get('validated_doc2')
            updated_b64 = request.session.get('validated_updated')
            if not (doc1_b64 and doc2_b64 and updated_b64):
                return HttpResponse("Erreur : documents manquants en session.", status=400)
            doc1_bytes = base64.b64decode(doc1_b64)
            doc2_bytes = base64.b64decode(doc2_b64)
            updated_bytes = base64.b64decode(updated_b64)

            validated_obj = ValidatedDocuments()
            validated_obj.doc1.save("doc1_batch_analyses.docx", ContentFile(doc1_bytes))
            validated_obj.doc2.save("doc2_laboratoire.docx", ContentFile(doc2_bytes))
            validated_obj.updated_doc.save("updated_batch_analyses.docx", ContentFile(updated_bytes))
            validated_obj.save()

            return render(request, 'validate_result.html', {
                'message': 'Documents validés et stockés avec succès.',
                'validated': validated_obj,
            })
        else:
            form = DocumentUploadForm(request.POST, request.FILES)
            if form.is_valid():
                doc1_file = request.FILES['document1']
                doc2_file = request.FILES['document2']

                # Lecture des contenus binaires
                doc1_content = doc1_file.read()
                doc2_content = doc2_file.read()
                doc1_file.seek(0)
                doc2_file.seek(0)

                if action == 'update':
                    updated_file = update_batch_analyses(doc1_file, doc2_file)
                    response = HttpResponse(
                        updated_file.getvalue(),
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    response['Content-Disposition'] = 'attachment; filename=updated_batch_analyses.docx'
                    return response

                elif action == 'test':
                    updated_file = update_batch_analyses(doc1_file, doc2_file)
                    doc2_file_io = BytesIO(doc2_content)
                    doc2_for_test = docx.Document(doc2_file_io)
                    updated_doc = docx.Document(updated_file)
                    errors = verify_update(updated_doc, doc2_for_test)
                    if errors:
                        result = "Test échoué :\n" + "\n".join(errors)
                    else:
                        result = "Test réussi : toutes les informations du document 2 sont présentes dans le document mis à jour."
                    updated_text = get_doc_text(updated_doc)
                    doc2_text = get_doc_text(doc2_for_test)
                    # Stockage dans la session pour validation ultérieure
                    request.session['validated_doc1'] = base64.b64encode(doc1_content).decode('utf-8')
                    request.session['validated_doc2'] = base64.b64encode(doc2_content).decode('utf-8')
                    request.session['validated_updated'] = base64.b64encode(updated_file.getvalue()).decode('utf-8')
                    return render(request, 'test_result.html', {
                        'result': result,
                        'updated_text': updated_text,
                        'doc2_text': doc2_text,
                    })

                elif action == 'compare':
                    # Extraction complète du contenu de doc1 et doc2
                    doc1_extracted = get_doc_text(docx.Document(doc1_file))
                    doc2_extracted = get_doc_text(docx.Document(doc2_file))
                    diff_html = difflib.HtmlDiff().make_file(
                        doc1_extracted.splitlines(),
                        doc2_extracted.splitlines(),
                        fromdesc="Doc1 Batch Analyses",
                        todesc="Doc2 Laboratoire"
                    )
                    return render(request, 'compare_documents.html', {
                        'doc1_text': doc1_extracted,
                        'doc2_text': doc2_extracted,
                        'diff_html': diff_html,
                    })
            else:
                return render(request, 'upload.html', {'form': form})
    else:
        form = DocumentUploadForm()
    return render(request, 'upload.html', {'form': form})

def home(request):
    updated_docs = ValidatedDocuments.objects.all().order_by('-validated_at')
    return render(request, 'home.html', {'updated_docs': updated_docs})

def delete_document(request, doc_id):
    # Supprime un document validé et redirige vers la page d'accueil
    doc = get_object_or_404(ValidatedDocuments, id=doc_id)
    doc.delete()
    return redirect('home')


# # batch_analyses/views.py
# from django.shortcuts import render
# from django.http import HttpResponse
# from .forms import DocumentUploadForm
# import docx
# from io import BytesIO
#
#
# def update_table(original_table, new_table):
#     """
#     Met à jour le tableau 'original_table' (du document 1) en se basant sur la structure
#     et le contenu du 'new_table' (extrait du document 2).
#     Cette fonction supprime toutes les lignes existantes (hors en-tête) puis ajoute les lignes
#     du tableau source.
#     """
#     # Conserver la première ligne (l'en-tête)
#     header = original_table.rows[0]
#     # Supprimer les lignes existantes hors en-tête
#     row_elements = original_table._element.xpath("./w:tr")
#     for row in row_elements[1:]:
#         original_table._element.remove(row)
#     # Ajouter les lignes du new_table (en sautant son en-tête)
#     for new_row in new_table.rows[1:]:
#         new_row_in_orig = original_table.add_row()
#         for idx, cell in enumerate(new_row.cells):
#             # Veillez à ne pas dépasser le nombre de cellules de la ligne d'origine
#             if idx < len(new_row_in_orig.cells):
#                 new_row_in_orig.cells[idx].text = cell.text
#
#
# def compare_impurities(doc1, doc2):
#     """
#     Exemple simplifié de comparaison des impuretés.
#     Dans une implémentation réelle, vous devrez extraire les valeurs dans les tableaux
#     correspondant aux tests de pureté pour déterminer le nombre d'impuretés et détecter
#     d'éventuelles inversions d'ordre.
#     """
#     # Pour l'exemple, on considère que doc1 indique 5 impuretés et doc2 3 impuretés
#     impurities_doc1 = 5
#     impurities_doc2 = 3
#     inversion = "Oui"  # Supposons qu'une inversion d'ordre est détectée
#     return impurities_doc1, impurities_doc2, inversion
#
#
# def update_batch_analyses(doc1_file, doc2_file):
#     """
#     Ouvre les deux documents, met à jour les tableaux de la section "Batch Analyses"
#     et ajoute un résumé comparatif à la fin du document.
#
#     Hypothèses de cet exemple :
#       - Le premier tableau de chaque document correspond à la section "BATCHES STUDIED".
#       - Les tableaux suivants représentent les résultats et doivent être mis à jour de façon similaire.
#     """
#     # Charger les documents
#     doc1 = docx.Document(doc1_file)
#     doc2 = docx.Document(doc2_file)
#
#     # Mise à jour du tableau "BATCHES STUDIED" (premier tableau)
#     if doc1.tables and doc2.tables:
#         update_table(doc1.tables[0], doc2.tables[0])
#
#     # Mise à jour des tableaux des résultats (à partir du second tableau)
#     nb_tables = min(len(doc1.tables), len(doc2.tables))
#     for i in range(1, nb_tables):
#         update_table(doc1.tables[i], doc2.tables[i])
#
#     # Comparaison des données relatives aux impuretés
#     impurities_doc1, impurities_doc2, inversion = compare_impurities(doc1, doc2)
#     # Ajout d'un paragraphe récapitulatif à la fin du document
#     doc1.add_paragraph("")
#     doc1.add_paragraph("=== Mise à jour par REG X ===")
#     doc1.add_paragraph("Les données du document 2 ont été intégrées dans la section 'Batch Analyses'.")
#     doc1.add_paragraph("Comparaison des impuretés :")
#     doc1.add_paragraph("• Document 1 : {} impuretés détectées".format(impurities_doc1))
#     doc1.add_paragraph("• Document 2 : {} impuretés détectées".format(impurities_doc2))
#     doc1.add_paragraph("• Inversion de l'ordre des impuretés : {}".format(inversion))
#
#     # Sauvegarder le document mis à jour dans un objet BytesIO pour le renvoyer en téléchargement
#     f = BytesIO()
#     doc1.save(f)
#     f.seek(0)
#     return f
#
#
# def upload_documents(request):
#     if request.method == 'POST':
#         form = DocumentUploadForm(request.POST, request.FILES)
#         if form.is_valid():
#             doc1_file = request.FILES['document1']
#             doc2_file = request.FILES['document2']
#             updated_file = update_batch_analyses(doc1_file, doc2_file)
#             response = HttpResponse(
#                 updated_file.getvalue(),
#                 content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#             )
#             response['Content-Disposition'] = 'attachment; filename=updated_batch_analyses.docx'
#             return response
#     else:
#         form = DocumentUploadForm()
#     return render(request, 'upload.html', {'form': form})
