# batch_analyses/tests.py

from django.test import TestCase, Client
from django.urls import reverse
from io import BytesIO
import docx

class UpdateBatchAnalysesTest(TestCase):
    def setUp(self):
        self.client = Client()

        # Création du document 1 (données de mars 2021)
        self.doc1_io = BytesIO()
        doc1 = docx.Document()
        doc1.add_heading("BATCH ANALYSES", level=1)
        doc1.add_heading("BATCHES STUDIED", level=2)
        table1 = doc1.add_table(rows=1, cols=5)
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = "Batch No."
        hdr_cells[1].text = "Date of manufacture"
        hdr_cells[2].text = "Batch size"
        hdr_cells[3].text = "Batch use"
        hdr_cells[4].text = "Manufacturing site"
        data1 = [
            ["2050717", "March 2021", "4,500,000", "Industrial batches", "LES LABORATOIRES"],
            ["2050957", "March 2021", "4,500,000", "Industrial batches", "LES LABORATOIRES"],
            ["2051204", "March 2021", "4,500,000", "Industrial batches", "LES LABORATOIRES"],
        ]
        for row in data1:
            row_cells = table1.add_row().cells
            for i, cell_val in enumerate(row):
                row_cells[i].text = cell_val

        doc1.add_heading("Résultats", level=2)
        table2 = doc1.add_table(rows=1, cols=4)
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].text = "Test"
        hdr_cells[1].text = "Batch 1"
        hdr_cells[2].text = "Batch 2"
        hdr_cells[3].text = "Batch 3"
        row_cells = table2.add_row().cells
        row_cells[0].text = "Impuretés détectées"
        row_cells[1].text = "5"
        row_cells[2].text = "5"
        row_cells[3].text = "5"

        doc1.save(self.doc1_io)
        self.doc1_io.seek(0)

        # Création du document 2 (données d'avril 2021)
        self.doc2_io = BytesIO()
        doc2 = docx.Document()
        doc2.add_heading("ANALYSES", level=1)
        doc2.add_heading("BATCHES STUDIED", level=2)
        table3 = doc2.add_table(rows=1, cols=5)
        hdr_cells = table3.rows[0].cells
        hdr_cells[0].text = "Batch No."
        hdr_cells[1].text = "Date of manufacture"
        hdr_cells[2].text = "Batch size"
        hdr_cells[3].text = "Batch use"
        hdr_cells[4].text = "Manufacturing site"
        data2 = [
            ["2051416", "April 2021", "2,000,000", "Industrial batches", "Laboratoires"],
            ["2051102", "April 2021", "2,000,000", "Industrial batches", "Laboratoires"],
            ["2051417", "April 2021", "2,000,000", "Industrial batches", "Laboratoires"],
        ]
        for row in data2:
            row_cells = table3.add_row().cells
            for i, cell_val in enumerate(row):
                row_cells[i].text = cell_val

        doc2.add_heading("Résultats", level=2)
        table4 = doc2.add_table(rows=1, cols=4)
        hdr_cells = table4.rows[0].cells
        hdr_cells[0].text = "Test"
        hdr_cells[1].text = "Batch 1"
        hdr_cells[2].text = "Batch 2"
        hdr_cells[3].text = "Batch 3"
        row_cells = table4.add_row().cells
        row_cells[0].text = "Impuretés détectées"
        row_cells[1].text = "3"
        row_cells[2].text = "3"
        row_cells[3].text = "3"
        doc2.save(self.doc2_io)
        self.doc2_io.seek(0)

    def test_update_document(self):
        # Envoi d'une requête POST vers l'URL associée à la vue upload_documents
        url = reverse('upload_documents')
        response = self.client.post(url, {
            'document1': self.doc1_io,
            'document2': self.doc2_io,
        })
        # Vérifier que la réponse est 200 et qu'un document DOCX est retourné
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # Charger le document généré
        updated_io = BytesIO(response.content)
        updated_doc = docx.Document(updated_io)

        # Vérifier la mise à jour du tableau "BATCHES STUDIED"
        # Le premier tableau doit désormais contenir les données du document 2
        updated_table = updated_doc.tables[0]
        # La première ligne est l'en-tête, la deuxième ligne doit être issue du document 2 (première ligne de données)
        first_data_row = updated_table.rows[1].cells
        self.assertEqual(first_data_row[0].text, "2051416")
        self.assertEqual(first_data_row[1].text, "April 2021")

        # Vérifier que le tableau "Résultats" a été mis à jour pour refléter "3" impuretés détectées
        updated_results_table = updated_doc.tables[1]
        results_row = updated_results_table.rows[1].cells
        self.assertEqual(results_row[1].text, "3")
        self.assertEqual(results_row[2].text, "3")
        self.assertEqual(results_row[3].text, "3")

        # Vérifier que le résumé de mise à jour est présent dans le document
        summary_exists = any("Mise à jour par REG X" in paragraph.text for paragraph in updated_doc.paragraphs)
        self.assertTrue(summary_exists)
